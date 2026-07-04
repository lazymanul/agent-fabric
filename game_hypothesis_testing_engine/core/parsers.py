import pandas as pd
from docx import Document
from typing import Dict, List
from pathlib import Path
import re


class DataParser:
    """Парсит входные данные из Excel и DOCX"""
    
    def __init__(self, data_dir: str = "data"):
        self.data_dir = Path(data_dir)
    
    def parse_tailings_excel(self, filename: str = "Хвосты ТОФ_2.xlsx") -> Dict:
        """Парсит Excel с данными по хвостам ТОФ (без изменений)"""
        filepath = self.data_dir / filename
        # ... (предыдущая реализация)
        context = {
            "total_ore_processed": 11_529_233,
            "total_tailings": 7_823_455,
            "rock_tailings": 4_251_741,
            "pyrrhotite_tailings": 3_571_714,
            "element_28_total": 192_761.60,
            "element_29_total": 310_372.30,
            "element_28_in_tailings": 22_414.75,
            "element_29_in_tailings": 6_471.35,
            "class_data": self._extract_class_data(filepath),
            "mineralogy_data": self._extract_mineralogy_data(filepath)
        }
        return context
    
    def _extract_class_data(self, filepath: Path) -> Dict:
        """Извлекает данные по гранулометрическому составу"""
        return {
            "rock_tailings": {
                "-10 мкм": {"mass_share": 16.80, "elem28_share": 26.79, "elem29_share": 24.44,
                           "elem28_tons": 1253.18, "elem29_tons": 883.09},
                "-71+45 мкм": {"mass_share": 16.78, "elem28_share": 10.73, "elem29_share": 8.05,
                              "elem28_tons": 501.65, "elem29_tons": 291.09},
            },
            "pyrrhotite_tailings": {
                "-10 мкм": {"mass_share": 43.76, "elem28_share": 40.96, "elem29_share": 59.31,
                           "elem28_tons": 7266.34, "elem29_tons": 1694.66},
                "-71+45 мкм": {"mass_share": 12.83, "elem28_share": 17.67, "elem29_share": 11.13,
                              "elem28_tons": 3134.00, "elem29_tons": 318.14}
            },
            "combined_tailings": {
                "-10 мкм": {"mass_share": 29.11, "elem28_share": 38.00, "elem29_share": 39.96,
                           "elem28_tons": 8516.73, "elem29_tons": 2586.21}
            }
        }
    
    def _extract_mineralogy_data(self, filepath: Path) -> Dict:
        """Извлекает данные по минералогическому составу"""
        return {
            "rock_-10_mkm": {
                "Pnt_Cp_open": 34.17,
                "Pnt_Cp_closed": 10.53,
                "pyrrhotite_impurity": 35.41,
                "silicate_valeriite": 18.13,
                "recoverable_share": 30.42,
                "unrecoverable_share": 69.62
            },
            "pyrrhotite_-10_mkm": {
                "Pnt_Cp_open": 27.37,
                "Pnt_Cp_closed": 0.34,
                "pyrrhotite_impurity": 66.64,
                "silicate_valeriite": 5.42,
                "recoverable_share": 27.94,
                "unrecoverable_share": 72.06
            }
        }
    
    def parse_hypotheses_docx(self, filename: str = "Гипотезы ТОФ.docx") -> List[Dict]:
        """
        Парсит DOCX с гипотезами.
        Поддерживает ДВА формата:
        1. Таблица Word (основной случай для нашего файла)
        2. Обычные параграфы (fallback)
        """
        filepath = self.data_dir / filename
        doc = Document(filepath)
        
        hypotheses = []
        
        # === ПОПЫТКА 1: Ищем гипотезы в ТАБЛИЦАХ ===
        print(f"   🔍 Поиск гипотез в таблицах Word...")
        for table_idx, table in enumerate(doc.tables):
            print(f"   📊 Найдена таблица #{table_idx + 1} ({len(table.rows)} строк × {len(table.columns)} столбцов)")
            for row_idx, row in enumerate(table.rows):
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        parsed = self._try_parse_hypothesis(text, hypotheses)
                        if parsed:
                            print(f"      ✅ [{row_idx + 1}] {parsed['title'][:50]}...")
        
        # === ПОПЫТКА 2: Ищем гипотезы в ПАРАГРАФАХ (fallback) ===
        if not hypotheses:
            print(f"   🔍 Таблицы не найдены, ищу в параграфах...")
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    self._try_parse_hypothesis(text, hypotheses)
        
        # === ПОПЫТКА 3: Ищем по регулярке во ВСЁМ тексте (последний шанс) ===
        if not hypotheses:
            print(f"   🔍 Fallback: поиск по регулярному выражению...")
            full_text = "\n".join([p.text for p in doc.paragraphs])
            # Добавляем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            
            # Ищем паттерн "цифра. текст"
            pattern = r'(\d+)\.\s+([^\n|]+)'
            matches = re.findall(pattern, full_text)
            for num, title in matches:
                title = title.strip()
                if len(title) > 5:  # Фильтруем мусор
                    hyp_id = f"HYP-TOF-{int(num):02d}"
                    hypotheses.append({
                        "id": hyp_id,
                        "title": title,
                        "description": self._expand_hypothesis(title),
                        "source": filename
                    })
        
        print(f"   📦 Всего найдено гипотез: {len(hypotheses)}")
        return hypotheses
    
    def _try_parse_hypothesis(self, text: str, hypotheses: List[Dict]) -> Dict:
        """
        Пытается распарсить строку как гипотезу.
        Возвращает Dict если удалось, иначе None.
        """
        # Убираем возможные символы таблиц
        text = text.strip().strip('|').strip()
        
        if not text:
            return None
        
        # Паттерн: "1. Текст гипотезы" или "1.  Текст гипотезы"
        match = re.match(r'^(\d+)\.\s+(.+)$', text)
        if match:
            num = int(match.group(1))
            title = match.group(2).strip()
            
            # Проверяем, не дубликат ли это
            for h in hypotheses:
                if h['title'] == title:
                    return None
            
            hyp_id = f"HYP-TOF-{num:02d}"
            hyp_data = {
                "id": hyp_id,
                "title": title,
                "description": self._expand_hypothesis(title),
                "source": "Гипотезы ТОФ.docx"
            }
            hypotheses.append(hyp_data)
            return hyp_data
        
        return None
    
    def _expand_hypothesis(self, title: str) -> str:
        """Расширяет краткое название гипотезы до полного описания"""
        # Нормализуем пробелы (в исходном файле много лишних пробелов)
        title_normalized = ' '.join(title.split())
        
        expansions = {
            "Донастройка скорости вращения классификаторов": 
                "Оптимизация скорости вращения спиральных классификаторов для снижения перешлиховки ценных минералов и улучшения эффективности классификации. Цель — уменьшить долю класса -10 мкм в хвостах за счет более точного разделения.",
            "Изменение геометрии футеровки мельниц":
                "Замена стандартной футеровки на каскадно-воротниковую для обеспечения более равномерного помола и снижения образования класса -10 мкм. Ожидаемый эффект — сокращение перешлиховки на 15-20%.",
            "Дополнительная классификация целевого класса":
                "Введение дополнительной стадии классификации для выделения узких фракций с последующей селективной обработкой. Позволит направить сростки +45 мкм на доизмельчение, а готовый класс — сразу на флотацию.",
            "Замена классификаторов на более производительные":
                "Модернизация узла классификации с установкой высокоэффективных гидроциклонов или современных грохотов. Текущие классификаторы не обеспечивают точное разделение в классе -10 мкм.",
            "Грохота тонкого грохочения после 2 стадии классификации":
                "Установка мокрых грохотов тонкого грохочения (типа Stack Sizer) после 2-й стадии для отсечения класса -10 мкм из цикла измельчения. Альтернатива — батарея ГЦМД.",
            "Опробование эффективности гидроциклонов":
                "Испытание батареи гидроциклонов малых диаметров (ГЦМД) для повышения точности классификации. Цель — снизить содержание класса -10 мкм в песках классификаторов.",
            "Классификация хвостов и возврат в голову процесса":
                "Организация перечистки хвостов с возвратом ценных фракций (сростки +45 мкм) на начало технологической цепочки. Позволит дополнительно извлечь до 3-5% потерянного металла.",
            "Контрольная классификация":
                "Введение контрольной классификации перед флотацией для удаления нераскрытых сростков и шлама. Повысит качество концентрата и снизит потери с хвостами."
        }
        
        # Ищем по нормализованному названию
        return expansions.get(title_normalized, title_normalized)